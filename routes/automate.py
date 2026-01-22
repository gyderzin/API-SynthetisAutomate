import os
import traceback
from io import BytesIO
from base64 import b64decode
import re
from fastapi import APIRouter, Depends, HTTPException
from sqlalchemy.orm import Session
from database import SessionLocal
from models.modelo import Modelo
from models.relatorio import Relatorio

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from uuid import uuid4
from datetime import datetime, timezone, timedelta

FUSO_BR = timezone(timedelta(hours=-4))  # Rondônia (UTC-4)


# =========================
# CONFIG
# =========================
BASE_DIR = r"C:\synthetis\relatórios"

router = APIRouter(
    prefix="/automate",
    tags=["Automate"]
)


# =========================
# DB SESSION
# =========================
def get_model_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


# =========================
# UTILITÁRIOS
# =========================
def is_base64_image(value: str) -> bool:
    return isinstance(value, str) and value.startswith("data:image") and ";base64," in value


def decode_base64_image(data_url: str) -> BytesIO:
    base64_str = re.sub(r"^data:image/\w+;base64,", "", data_url)
    image_bytes = b64decode(base64_str)
    return BytesIO(image_bytes)


def clear_paragraph(paragraph):
    for run in paragraph.runs:
        paragraph._element.remove(run._element)


# ======================================
# SUBSTITUIÇÃO DE TEXTOS
# ======================================
def substituir_textos_no_documento(doc: Document, dados: dict):
    for paragraph in doc.paragraphs:
        for key, value in dados.items():

            # NÃO SUBSTITUI SE VALUE ESTIVER VAZIO / NONE / ""
            if value in [None, "", " ", "null"]:
                continue

            if is_base64_image(value):
                continue

            placeholder = f"{{{{{key}}}}}"

            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(value))
                paragraph.style.font.size = Pt(10)
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in dados.items():

                        if value in [None, "", " ", "null"]:
                            continue

                        if is_base64_image(value):
                            continue

                        placeholder = f"{{{{{key}}}}}"

                        if placeholder in paragraph.text:
                            paragraph.text = paragraph.text.replace(placeholder, str(value))
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in paragraph.runs:
                                run.font.size = Pt(10)



# ======================================
# CABEÇALHO / RODAPÉ
# ======================================
def substituir_textos_em_cabecalho_rodape(doc: Document, dados: dict):
    for section in doc.sections:
        for area in [section.header, section.footer]:

            for paragraph in area.paragraphs:
                for key, value in dados.items():
                    if is_base64_image(value):
                        continue

                    placeholder = f"{{{{{key}}}}}"

                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, str(value))
                        paragraph.style.font.size = Pt(10)
                        for run in paragraph.runs:
                            run.font.size = Pt(10)

            for table in area.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for key, value in dados.items():
                                if is_base64_image(value):
                                    continue

                                placeholder = f"{{{{{key}}}}}"

                                if placeholder in paragraph.text:
                                    paragraph.text = paragraph.text.replace(placeholder, str(value))
                                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    for run in paragraph.runs:
                                        run.font.size = Pt(10)


# ======================================
# SUBSTITUIÇÃO DE IMAGENS
# ======================================
def substituir_imagens_no_documento(doc: Document, dados: dict):
    for paragraph in doc.paragraphs:
        for key, value in dados.items():

            # Se não houver imagem, NÃO substitui e NÃO remove o placeholder
            if not value or value in ["", " ", None, "null"]:
                continue

            placeholder = f"{{{{{key}}}}}"

            # Só substitui se o valor realmente for uma imagem base64
            if placeholder in paragraph.text and is_base64_image(value):
                clear_paragraph(paragraph)
                run = paragraph.add_run()
                run.add_picture(decode_base64_image(value), height=Cm(10))

    # Agora nas tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in dados.items():

                        # Se não houver imagem válida, pula
                        if not value or value in ["", " ", None, "null"]:
                            continue

                        placeholder = f"{{{{{key}}}}}"

                        if placeholder in paragraph.text and is_base64_image(value):
                            clear_paragraph(paragraph)
                            run = paragraph.add_run()
                            run.add_picture(decode_base64_image(value), height=Cm(10))



# ======================================
# INSERIR PENDÊNCIAS
# ======================================
def inserir_pendencias_no_documento(doc: Document, pendencias: list[dict], chavePendencia: str):
    placeholder = f"{{{{{chavePendencia}}}}}"

    def process_paragraph(paragraph):
        if placeholder in paragraph.text:
            clear_paragraph(paragraph)

            if not pendencias:
                paragraph.add_run("Nenhuma pendência relatada.\n")
                return True

            for idx, pendencia in enumerate(pendencias, 1):
                titulo = pendencia.get("titulo", "Sem título")
                descricao = pendencia.get("descriçao") or pendencia.get("descricao") or ""
                imagem = pendencia.get("imagem")

                # título
                par_titulo = paragraph.insert_paragraph_before(f"{idx}. {titulo}")
                par_titulo.runs[0].bold = True
                par_titulo.runs[0].font.size = Pt(10)

                # descrição
                if descricao:
                    par_desc = paragraph.insert_paragraph_before(descricao)
                    par_desc.runs[0].font.size = Pt(10)

                # imagem
                if imagem and is_base64_image(imagem):
                    par_img = paragraph.insert_paragraph_before()
                    run_img = par_img.add_run()
                    run_img.add_picture(decode_base64_image(imagem), width=Cm(10))

            return True

        return False

    # busca no corpo
    for paragraph in doc.paragraphs:
        if process_paragraph(paragraph):
            return

    # busca em tabelas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if process_paragraph(paragraph):
                        return


# ======================================
# ROTA PRINCIPAL (armazenando em DISCO)
# ======================================
@router.post("/gerar-doc") 
def gerar_documento(payload: dict, db: Session = Depends(get_model_db)):
    try:
        import json

        modelo_id = payload.get("modelo_id")
        dados = payload.get("dados", {})
        pendencias = payload.get("pendencias", [])
        chavePendencia = payload.get("chavePendencia")
        responsavel = payload.get("responsavel")
        equipamento = str(payload.get("equipamento", "")).strip()

        # <-- Nome vindo do front-end
        nome_relatorio = payload.get("nome_relatorio")

        # NOME EXATO que o front-end envia
        itens_pendentes_raw = payload.get("itens_pendentes")

        # Salvar como JSON string
        itens_pendentes = json.dumps(itens_pendentes_raw)

        modelo = db.query(Modelo).filter(Modelo.id == modelo_id).first()
        if not modelo or not modelo.documento_modelo:
            raise HTTPException(status_code=404, detail="Modelo não encontrado ou sem documento")

        doc = Document(BytesIO(modelo.documento_modelo))

        substituir_textos_no_documento(doc, dados)
        substituir_textos_em_cabecalho_rodape(doc, dados)
        substituir_imagens_no_documento(doc, dados)

        if chavePendencia:
            inserir_pendencias_no_documento(doc, pendencias, chavePendencia)

        # ============================
        #  DEFINIR NOME DO ARQUIVO
        # ============================

        if nome_relatorio and str(nome_relatorio).strip() != "":
            nome_arquivo_original = nome_relatorio.strip()
        else:
            nome_arquivo_original = f"{modelo.titulo.strip()} {equipamento}".strip()

        # Sempre gerar nome físico temporário
        _, ext = os.path.splitext("arquivo.docx")
        nome_arquivo_fisico = f"{uuid4()}{ext}"
        caminho_final = os.path.join(BASE_DIR, nome_arquivo_fisico)

        doc.save(caminho_final)

        novo = Relatorio(
            modelo=modelo.titulo,
            emissor=responsavel,
            equipe=modelo.equipe,
            nome_arquivo=nome_arquivo_original,  
            caminho_arquivo=caminho_final,
            item_pendente=itens_pendentes
        )

        db.add(novo)
        db.commit()

        return {"status": "ok", "mensagem": "Relatório armazenado no disco com sucesso"}

    except Exception as e:
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Erro ao gerar documento: {str(e)}")



@router.post("/resolver-itens-pendentes")
def resolver_pendencias(payload: dict, db: Session = Depends(get_model_db)):
    try:
        relatorio_id = payload.get("relatorio_id")
        imagens = payload.get("imagens", {})

        if not relatorio_id:
            raise HTTPException(status_code=400, detail="relatorio_id é obrigatório.")

        relatorio = db.query(Relatorio).filter(Relatorio.id == relatorio_id).first()
        if not relatorio:
            raise HTTPException(status_code=404, detail="Relatório não encontrado.")

        if not os.path.exists(relatorio.caminho_arquivo):
            raise HTTPException(status_code=404, detail="Arquivo do relatório não encontrado.")

        doc = Document(relatorio.caminho_arquivo)

        def substituir_imagem_placeholder(paragraph, chave, base64img):
            placeholder = f"{{{{{chave}}}}}"
            if placeholder in paragraph.text and is_base64_image(base64img):
                clear_paragraph(paragraph)
                run = paragraph.add_run()
                run.add_picture(decode_base64_image(base64img), height=Cm(10))
                return True
            return False

        for paragraph in doc.paragraphs:
         for chave, img in imagens.items():
             if substituir_imagem_placeholder(paragraph, chave, img):
                 break  # <- interrompe para não tentar substituir de novo


        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for chave, img in imagens.items():
                            if substituir_imagem_placeholder(paragraph, chave, img):
                                break

        nome_final = f"{uuid4()}"
        caminho_final = os.path.join(BASE_DIR, nome_final)

        doc.save(caminho_final)

        # --- Atualizações no banco ---
        relatorio.caminho_arquivo = caminho_final

        # Zerar pendências com lista vazia
        relatorio.item_pendente = None

        # Atualizar data com horário local (UTC-4)
        relatorio.emitido_em = datetime.now(FUSO_BR)

        db.commit()

        return {
            "status": "ok",
            "mensagem": "Pendências resolvidas e relatório atualizado.",
            "novo_arquivo": caminho_final
        }

    except Exception as e:
        print(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"Erro ao resolver pendências: {str(e)}")



